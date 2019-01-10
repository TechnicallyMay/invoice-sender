from docx import Document
import translate


class Invoice():

    def __init__(self, customer, owner):
        self.customer = customer
        self.owner = owner
        name = self.customer.data["Name"].replace(" ", "").lower()
        self.file_name = "../data/temp/%d%s.docx" % (self.customer.index, name)
        self.template = Document("../data/invoice_template.docx")
        self.add_charges(self.customer.charges)
        self.replace_key_words()
        self.template.save(self.file_name)


    def replace_key_words(self):
        for paragraph in self.template.paragraphs:
            paragraph.text = self.translate(paragraph.text)
        for table in self.template.tables:
            for i in range(len(table.rows)):
                for cell in table.row_cells(i):
                    for paragraph in cell.paragraphs:
                        paragraph.text = self.translate(paragraph.text)


    def translate(self, text):
        text = translate.translate(text, self.owner.data, self.customer.data)
        return text


    def add_charges(self, charge_dict):
        charge_table = self.template.tables[1]
        charges = charge_dict.copy()
        for i in range(1, len(charge_table.rows) - 1):
            try:
                charge = charges.pop(0)
                charge_total = float(charge[0]) * float(charge[2])
                charge.append(charge_total)
            except IndexError:
                charge = ("", "", "", "")
            cells = charge_table.row_cells(i)
            qty = str(charge[0])
            description = charge[1]
            unit_price = charge[2]
            total_charge = charge[3]
            if unit_price != "":
                unit_price = "${:0.2f}".format(unit_price)
                total_charge = "${:0.2f}".format(total_charge)
            cells[0].paragraphs[0].text = qty
            cells[1].paragraphs[0].text = description
            cells[2].paragraphs[0].text = unit_price
            cells[3].paragraphs[0].text = total_charge
        final_row = charge_table.row_cells(len(charge_table.rows) - 1)
        final_row[3].paragraphs[0].text = "${:0.2f}".format(self.customer.total)
